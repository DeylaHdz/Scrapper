import io
import pdfplumber
from docx import Document
from bs4 import BeautifulSoup
from pathlib import Path
import requests
from utils import clean_text
from playwright.sync_api import sync_playwright


USER_AGENT = "SimpleDocScraper/1.0 (+for demo; contact admin@example.com)"
ALLOWED_EXTS = {".pdf", ".docx", ".txt"}

def extract_pdf(file_bytes: bytes):
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            text_parts.append(txt)
    return clean_text("\n\n".join(text_parts)), {"page_count": len(pdf.pages)}

def extract_docx(file_bytes: bytes):
    doc = Document(io.BytesIO(file_bytes))
    paras = [p.text for p in doc.paragraphs]
    return clean_text("\n".join(paras)), {"paragraphs": len(doc.paragraphs)}

def extract_txt(file_bytes: bytes):
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore")
    return clean_text(text), {}

def extract_html(html_bytes: bytes, url: str = ""):
    soup = BeautifulSoup(html_bytes, "html.parser")
    title = soup.title.string.strip() if soup.title else ""
    for t in soup(["script", "style", "noscript"]):
        t.decompose()
    text = clean_text(soup.get_text(separator="\n"))
    return text, {"title": title, "source_url": url}

def extract_html_js(url: str):
    """
    Extrae texto renderizado con JS en un sitio SPA (React/Vite).
    """
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto(url, timeout=60000)

        # Espera hasta que React termine de hidratar y renderizar
        page.wait_for_load_state("networkidle")

        # Extrae el título
        title = page.title()

        # Extrae todo el texto visible en el body
        text = page.inner_text("body")

        browser.close()

    return clean_text(text), {"title": title, "source_url": url}

def extract_from_file(filename: str, file_bytes: bytes):
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(file_bytes)
    elif ext == ".docx":
        return extract_docx(file_bytes)
    elif ext == ".txt":
        return extract_txt(file_bytes)
    else:
        raise ValueError(f"Extensión no permitida: {ext}")

def extract_from_url(url: str):
    ext = Path(url.split("?")[0].split("#")[0]).suffix.lower()

    # Si es HTML → usar Playwright siempre
    if ext in {".htm", ".html", ""} or url.startswith("http"):
        return extract_html_js(url)

    # Archivos tradicionales
    headers = {"User-Agent": USER_AGENT}
    r = requests.get(url, headers=headers, timeout=30)
    r.raise_for_status()
    content_type = r.headers.get("Content-Type", "").lower()
    content = r.content

    if "pdf" in content_type or ext == ".pdf":
        return extract_pdf(content)
    if "word" in content_type or ext == ".docx":
        return extract_docx(content)
    if "text/plain" in content_type or ext == ".txt":
        return extract_txt(content)

    raise ValueError(f"No se pudo interpretar el tipo de contenido: {content_type or ext}")