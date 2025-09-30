from flask import Flask, request, jsonify, render_template_string
import io
import re
import tempfile
import requests
from pathlib import Path

# PDF
import pdfplumber
# DOCX
from docx import Document
# HTML
from bs4 import BeautifulSoup

ALLOWED_EXTS = {".pdf", ".docx", ".txt"}
USER_AGENT = "SimpleDocScraper/1.0 (+for demo; contact admin@example.com)"

app = Flask(__name__)

def clean_text(s: str) -> str:
    # Normaliza espacios y saltos de línea
    s = re.sub(r"\r\n|\r", "\n", s)
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

def guess_headings(text: str, min_len=3, max_len=120):
    headings = []
    for line in text.splitlines():
        ln = line.strip()
        if not (min_len <= len(ln) <= max_len):
            continue
        if (ln.isupper() and any(c.isalpha() for c in ln)) \
           or re.match(r"^(Cap[ií]tulo|Secci[oó]n|T[ií]tulo|Resumen|Abstract|Introducci[oó]n)\b", ln, re.I) \
           or re.match(r"^(\d+(\.\d+)*\s+).{2,}$", ln):
            headings.append(ln)
    seen = set()
    uniq = []
    for h in headings:
        if h.lower() not in seen:
            uniq.append(h)
            seen.add(h.lower())
    return uniq[:100]

def extract_pdf(file_bytes: bytes):
    meta = {}
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        meta["page_count"] = len(pdf.pages)
        for page in pdf.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            text_parts.append(txt)
    text = "\n\n".join(text_parts)
    return clean_text(text), meta

def extract_docx(file_bytes: bytes):
    f = io.BytesIO(file_bytes)
    doc = Document(f)
    paras = []
    for p in doc.paragraphs:
        paras.append(p.text)
    text = "\n".join(paras)
    meta = {"paragraphs": len(doc.paragraphs)}
    return clean_text(text), meta

def extract_txt(file_bytes: bytes):
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="ignore")
    return clean_text(text), {}

def extract_html(html_bytes: bytes, url: str = ""):
    soup = BeautifulSoup(html_bytes, "html.parser")
    title = (soup.title.string.strip() if soup.title and soup.title.string else "")
    for t in soup(["script", "style", "noscript"]):
        t.decompose()
    text = soup.get_text(separator="\n")
    text = clean_text(text)
    meta = {"title": title, "source_url": url}
    return text, meta

def fetch_url(url: str) -> tuple[str, dict]:
    headers = {"User-Agent": USER_AGENT}
    r = requests.get(url, headers=headers, timeout=30)
    r.raise_for_status()
    content_type = r.headers.get("Content-Type", "").lower()
    content = r.content

    ext = Path(url.split("?")[0].split("#")[0]).suffix.lower()

    if "text/html" in content_type or ext in {".htm", ".html"}:
        return extract_html(content, url)

    if "pdf" in content_type or ext == ".pdf":
        text, meta = extract_pdf(content)
        meta.update({"source_url": url, "content_type": "application/pdf"})
        return text, meta

    if "word" in content_type or "officedocument.wordprocessingml" in content_type or ext == ".docx":
        text, meta = extract_docx(content)
        meta.update({"source_url": url, "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
        return text, meta

    if "text/plain" in content_type or ext == ".txt":
        text, meta = extract_txt(content)
        meta.update({"source_url": url, "content_type": "text/plain"})
        return text, meta
    try:
        text, meta = extract_pdf(content)
        meta.update({"source_url": url, "content_type": content_type or "unknown"})
        if text.strip():
            return text, meta
    except Exception:
        pass

    raise ValueError(f"No se pudo interpretar el tipo de contenido: {content_type or ext or 'desconocido'}")

def limit_text(text: str, limit: int | None):
    if limit is None or limit <= 0:
        return text
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[... truncado ...]"

HOME_HTML = """
<!doctype html>
<html lang="es">
    <head>
        <meta charset="utf-8"/>
        <meta name="viewport" content="width=device-width, initial-scale=1"/>
        <title>Scraper</title>
        <script type="module" src="https://unpkg.com/ionicons@7.1.0/dist/ionicons/ionicons.esm.js"></script>
        <script nomodule src="https://unpkg.com/ionicons@7.1.0/dist/ionicons/ionicons.js"></script>
        <style>
            body { font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; max-width: 820px; margin: 2rem auto; padding: 0 1rem; }
            .card { border: 1px solid #ddd; border-radius: 14px; padding: 1rem 1.2rem; margin-bottom: 1rem; }
            h1 { font-size: 1.6rem; }
            label { display: block; font-weight: 600; margin-top: .5rem; }
            input[type="text"]{ width: 100%; padding: .6rem; border: 1px solid #ccc; border-radius: 10px; }
            input[type="file"]{ margin-top: .4rem; }
            button { margin-top: .8rem; padding: .6rem 1rem; border-radius: 10px; border: 1px solid #555; background: #111; color: #fff; cursor: pointer;}
            pre { background: #111; color: #e3e3e3; padding: 1rem; border-radius: 10px; overflow:auto; }
            small { color: #666; }
        </style>
    </head>
    <body>
        <h1>Scraper de Documentos & URLs</h1>

        <div class="card">
            <h3>Subir archivo</h3>
            <form id="fileForm">
            <label>Archivo (.pdf, .docx, .txt):</label>
            <input name="file" type="file" required />
            <label>Límite de caracteres del texto (opcional):</label>
            <input name="limit" type="text" placeholder="p. ej. 5000" />
            <button type="submit">Extraer</button>
            </form>
        </div>

        <div class="card">
            <h3>Desde URL</h3>
            <form id="urlForm">
            <label>URL:</label>
            <input name="url" type="text" placeholder="https://ejemplo.com/documento.pdf" required />
            <label>Límite de caracteres del texto (opcional):</label>
            <input name="limit" type="text" placeholder="p. ej. 5000" />
            <button type="submit">Extraer</button>
            </form>
        </div>

        <div class="card">
            <h3>Salida</h3>
            <pre id="out">—</pre>
            <small>Tip: la API también está en <code>POST /extract</code>.</small>
        </div>

        <script>
        async function toJSON(form, endpoint="/extract") {
        const fd = new FormData(form);
        const resp = await fetch(endpoint, { method:"POST", body: fd });
        const data = await resp.json();
        return data;
        }
        document.getElementById("fileForm").addEventListener("submit", async (e) => {
        e.preventDefault();
        const out = document.getElementById("out");
        out.textContent = "Procesando archivo...";
        try {
            const data = await toJSON(e.target);
            out.textContent = JSON.stringify(data, null, 2);
        } catch (err) {
            out.textContent = "Error: " + err;
        }
        });
        document.getElementById("urlForm").addEventListener("submit", async (e) => {
        e.preventDefault();
        const out = document.getElementById("out");
        out.textContent = "Descargando URL...";
        try {
            const data = await toJSON(e.target);
            out.textContent = JSON.stringify(data, null, 2);
        } catch (err) {
            out.textContent = "Error: " + err;
        }
        });
        </script>
    </body> 
</html>
"""

@app.get("/")
def home():
    return render_template_string(HOME_HTML)

@app.post("/extract")
def extract():
    try:
        limit_param = request.form.get("limit") or request.args.get("limit")
        limit = int(limit_param) if limit_param else None

        # Prioridad: archivo subido
        if "file" in request.files and request.files["file"].filename:
            file = request.files["file"]
            filename = file.filename
            ext = Path(filename).suffix.lower()

            if ext not in ALLOWED_EXTS:
                return jsonify({"ok": False, "error": f"Extensión no permitida: {ext}"}), 400

            file_bytes = file.read()
            if ext == ".pdf":
                text, meta = extract_pdf(file_bytes)
            elif ext == ".docx":
                text, meta = extract_docx(file_bytes)
            else:
                text, meta = extract_txt(file_bytes)

            headings = guess_headings(text)
            resp = {
                "ok": True,
                "source": "upload",
                "filename": filename,
                "metadata": meta | {"word_count": len(text.split())},
                "headings": headings,
                "text": limit_text(text, limit),
            }
            return jsonify(resp), 200

        # Alternativa: URL
        url = request.form.get("url") or request.args.get("url")
        if url:
            text, meta = fetch_url(url)
            headings = guess_headings(text)
            resp = {
                "ok": True,
                "source": "url",
                "url": url,
                "metadata": meta | {"word_count": len(text.split())},
                "headings": headings,
                "text": limit_text(text, limit),
            }
            return jsonify(resp), 200

        return jsonify({"ok": False, "error": "Proporciona un archivo o una URL."}), 400

    except requests.RequestException as e:
        return jsonify({"ok": False, "error": f"Error HTTP al descargar: {e}"}), 502
    except Exception as e:
        return jsonify({"ok": False, "error": f"{type(e).__name__}: {e}"}), 500

if __name__ == "__main__":
    app.run(debug=True, port=5000)
